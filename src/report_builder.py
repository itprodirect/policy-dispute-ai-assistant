from __future__ import annotations

import argparse
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Iterable

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from rich import print
from .schemas import DisputeReport


@dataclass
class PolicyReport:
    policy_name: str
    source_path: str
    num_sections: int
    num_unknown_sections: int
    num_meta_sections: int
    sections_substantive: List[Dict[str, Any]]
    sections_meta: List[Dict[str, Any]]


def build_policy_report(summary_data: Dict[str, Any]) -> PolicyReport:
    """Normalize raw JSON from the summarizer into a PolicyReport object."""
    policy_name = summary_data.get("policy_id") or "<unknown_policy>"
    source_path = summary_data.get("policy_path") or ""

    sections: List[Dict[str, Any]] = summary_data.get("sections", [])
    num_sections = len(sections)

    num_unknown_sections = 0
    sections_meta: List[Dict[str, Any]] = []
    sections_substantive: List[Dict[str, Any]] = []

    for s in sections:
        name = (s.get("section_name") or "").strip()
        role = (s.get("section_role") or "substantive").lower()

        if name.upper().startswith("UNKNOWN"):
            num_unknown_sections += 1

        if role == "meta":
            sections_meta.append(s)
        else:
            sections_substantive.append(s)

    num_meta_sections = len(sections_meta)

    return PolicyReport(
        policy_name=policy_name,
        source_path=source_path,
        num_sections=num_sections,
        num_unknown_sections=num_unknown_sections,
        num_meta_sections=num_meta_sections,
        sections_substantive=sections_substantive,
        sections_meta=sections_meta,
    )


def _render_list_block(title: str, items: List[str]) -> str:
    if not items:
        return ""
    lines = [f"**{title}**"]
    for item in items:
        lines.append(f"- {item}")
    lines.append("")  # blank line after block
    return "\n".join(lines)


def render_section_markdown(section: Dict[str, Any]) -> str:
    """Render a single section (already filtered to substantive) to Markdown."""
    name = section.get("section_name") or "UNKNOWN SECTION"
    summary = section.get("summary_overall") or ""

    key_cov = section.get("key_coverages") or []
    key_exc = section.get("key_exclusions") or []
    conds = section.get("conditions_notable") or []
    disputes = section.get("dispute_angles_possible") or []

    parts: List[str] = []
    parts.append(f"### {name}")
    parts.append("")
    if summary:
        parts.append(summary.strip())
        parts.append("")

    parts.append(_render_list_block("Key coverages", key_cov))
    parts.append(_render_list_block("Key exclusions", key_exc))
    parts.append(_render_list_block("Notable conditions / duties", conds))
    parts.append(_render_list_block("Potential dispute angles", disputes))

    # Remove any empty chunks
    parts = [p for p in parts if p.strip() != ""]
    parts.append("")  # trailing blank line for spacing
    return "\n".join(parts)


def render_markdown(report: PolicyReport) -> str:
    """Render the whole policy report to Markdown."""
    header_lines = [
        f"# Policy summary – {report.policy_name}",
        "",
        f"- Source file: `{report.source_path}`" if report.source_path else "",
        f"- Total sections: {report.num_sections}",
        f"- UNKNOWN sections: {report.num_unknown_sections}",
        f"- Meta / non-substantive sections: {report.num_meta_sections}",
        "",
        "> **Framing:** This is an AI-generated summary of policy language only.",
        "> It is *not* legal advice and does not create coverage or rights.",
        "",
    ]

    header_lines = [l for l in header_lines if l != ""]

    body_lines: List[str] = []
    body_lines.extend(header_lines)

    # Substantive / claim-facing sections
    body_lines.append("## Substantive sections (claim-facing)")
    body_lines.append("")
    if not report.sections_substantive:
        body_lines.append("_No substantive sections detected._")
        body_lines.append("")
    else:
        for s in report.sections_substantive:
            body_lines.append(render_section_markdown(s))

    # Meta sections appendix
    if report.sections_meta:
        body_lines.append(
            "## Meta / regulatory / boilerplate sections (not analyzed for coverage)"
        )
        body_lines.append("")
        for s in report.sections_meta:
            name = s.get("section_name") or "UNKNOWN"
            body_lines.append(f"- {name}")
        body_lines.append("")

    return "\n".join(body_lines)


def generate_report_for_summary(summary_json_path: Path) -> Path:
    """Load a single <policy>.json and emit a <policy>.report.md file next to it."""
    print(f"[bold]Building report from:[/bold] {summary_json_path}")
    data = json.loads(summary_json_path.read_text(encoding="utf-8"))
    report = build_policy_report(data)
    markdown = render_markdown(report)

    md_path = summary_json_path.with_suffix(".report.md")
    md_path.write_text(markdown, encoding="utf-8")
    print(f"[green]Wrote report:[/green] {md_path}")

    if report.num_sections and report.num_unknown_sections / report.num_sections > 0.3:
        print(
            f"[yellow][WARN][/yellow] {report.num_unknown_sections}/{report.num_sections} sections "
            "are UNKNOWN. Section detection may need tuning."
        )

    return md_path


def iter_input_files(paths: Iterable[str], pattern: str = "*.json") -> Iterable[Path]:
    """
    Expand CLI inputs into actual JSON files.

    - File path -> yield as-is
    - Directory  -> glob for pattern (default *.json)
    - Nonexistent path -> warn and skip
    """
    for raw in paths:
        p = Path(raw)
        if p.is_file():
            if p.suffix.lower() == ".json":
                yield p
        elif p.is_dir():
            for json_path in sorted(p.glob(pattern)):
                if json_path.is_file() and json_path.suffix.lower() == ".json":
                    yield json_path
        else:
            print(
                f"[yellow]Warning:[/yellow] path does not exist, skipping: {raw}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build claim-facing Markdown reports from policy summary JSON files."
    )
    parser.add_argument(
        "paths",
        nargs="+",
        help="One or more summary JSON files and/or directories containing JSON files.",
    )
    parser.add_argument(
        "--pattern",
        default="*.json",
        help="Glob pattern when expanding directories (default: *.json).",
    )
    parser.add_argument(
        "--skip-existing",
        action="store_true",
        help="Skip any JSON where the corresponding .report.md already exists.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    files = list(iter_input_files(args.paths, args.pattern))

    if not files:
        print("[red]No summary JSON files found. Nothing to do.[/red]")
        return

    print(f"[bold]Found {len(files)} summary JSON file(s) to process.[/bold]")

    count = 0
    for json_path in files:
        out_path = json_path.with_suffix(".report.md")
        if args.skip_existing and out_path.exists():
            print(
                f"[yellow]Skipping (report already exists):[/yellow] {json_path}")
            continue
        try:
            generate_report_for_summary(json_path)
            count += 1
        except Exception as e:
            print(f"[red]Error building report for {json_path}:[/red] {e}")

    print(f"[bold]Done.[/bold] Built {count} report file(s).")


def render_dispute_markdown(report: DisputeReport) -> str:
    """
    Render a DisputeReport (A–G structure) to Markdown.
    """
    lines: List[str] = []

    # Header
    title = "Policy dispute summary"
    lines.append(f"# {title}")
    lines.append("")

    meta_lines: List[str] = []
    if report.policy_id:
        meta_lines.append(f"- Policy: `{report.policy_id}`")
    if report.denial_id:
        meta_lines.append(f"- Denial: `{report.denial_id}`")

    if meta_lines:
        lines.extend(meta_lines)
        lines.append("")

    lines.append(
        "> **Framing:** This is an AI-generated analysis of policy language and a denial letter.\n"
        "> It is *not* legal advice and does not create coverage, rights, or an attorney–client relationship."
    )
    lines.append("")

    # A. Plain-English overview
    lines.append("## A. Plain-English overview of the dispute")
    lines.append("")
    if report.plain_summary.strip():
        lines.append(report.plain_summary.strip())
    else:
        lines.append("_No overview available._")
    lines.append("")

    # Helper for B/C/D
    def _append_points_section(heading: str, points: List[Any]) -> None:
        lines.append(heading)
        lines.append("")
        if not points:
            lines.append("_None identified._")
        else:
            for p in points:
                text = getattr(p, "text", "").strip()
                if not text:
                    continue
                citation = getattr(p, "citation", None)
                if citation:
                    lines.append(f"- {text} ({citation})")
                else:
                    lines.append(f"- {text}")
        lines.append("")

    _append_points_section(
        "## B. Coverage highlights that may support the insured",
        report.coverage_highlights,
    )
    _append_points_section(
        "## C. Key exclusions / limitations that may hurt the insured",
        report.exclusions_limitations,
    )
    _append_points_section(
        "## D. Denial reasons (as stated or implied by the insurer)",
        report.denial_reasons,
    )

    # E. Dispute angles
    lines.append("## E. Possible dispute angles to explore (not legal advice)")
    lines.append("")
    if not report.dispute_angles:
        lines.append("_No dispute angles identified._")
    else:
        for angle in report.dispute_angles:
            text = angle.text.strip()
            if not text:
                continue
            if angle.citations:
                cits = ", ".join(c for c in angle.citations if c.strip())
                if cits:
                    lines.append(f"- {text}  \n  _Citations: {cits}_")
                else:
                    lines.append(f"- {text}")
            else:
                lines.append(f"- {text}")
    lines.append("")

    # F. Missing info
    lines.append("## F. Missing information / suggested next steps")
    lines.append("")
    if not report.missing_info:
        lines.append("_No specific missing information identified._")
    else:
        for item in report.missing_info:
            s = str(item).strip()
            if s:
                lines.append(f"- {s}")
    lines.append("")

    # G. Confidence & verify clauses
    lines.append("## G. Confidence and clauses to verify")
    lines.append("")
    conf = report.confidence
    if conf.score is not None:
        lines.append(f"- **Confidence score (0–1):** {conf.score:.2f}")
    if conf.notes.strip():
        lines.append(f"- **Notes:** {conf.notes.strip()}")
    if conf.verify_clauses:
        lines.append("- **Clauses / sections to double-check:**")
        for clause in conf.verify_clauses:
            s = str(clause).strip()
            if s:
                lines.append(f"  - {s}")
    if not any(
        [
            conf.score is not None,
            conf.notes.strip(),
            bool(conf.verify_clauses),
        ]
    ):
        lines.append("_No explicit confidence metadata provided._")
    lines.append("")

    return "\n".join(lines)


def render_dispute_docx(report: DisputeReport) -> bytes:
    """
    Render a DisputeReport (A–G structure) to a Word document (.docx).

    Returns the document as bytes suitable for download.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Policy Dispute Summary", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata
    if report.policy_id or report.denial_id:
        meta_para = doc.add_paragraph()
        if report.policy_id:
            meta_para.add_run(f"Policy: {report.policy_id}").bold = True
            if report.denial_id:
                meta_para.add_run("  |  ")
        if report.denial_id:
            meta_para.add_run(f"Denial: {report.denial_id}").bold = True

    # Framing disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer.add_run("Framing: ").bold = True
    disclaimer.add_run(
        "This is an AI-generated analysis of policy language and a denial letter. "
        "It is not legal advice and does not create coverage, rights, or an "
        "attorney–client relationship."
    ).italic = True

    doc.add_paragraph()  # spacing

    # A. Plain-English overview
    doc.add_heading("A. Plain-English Overview of the Dispute", level=1)
    if report.plain_summary.strip():
        doc.add_paragraph(report.plain_summary.strip())
    else:
        doc.add_paragraph("No overview available.").italic = True

    # Helper for B/C/D sections with points
    def _add_points_section(heading: str, points: List[Any]) -> None:
        doc.add_heading(heading, level=1)
        if not points:
            p = doc.add_paragraph("None identified.")
            p.runs[0].italic = True
        else:
            for pt in points:
                text = getattr(pt, "text", "").strip()
                if not text:
                    continue
                citation = getattr(pt, "citation", None)
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(text)
                if citation:
                    para.add_run(f" ({citation})").italic = True

    _add_points_section(
        "B. Coverage Highlights That May Support the Insured",
        report.coverage_highlights,
    )
    _add_points_section(
        "C. Key Exclusions / Limitations That May Hurt the Insured",
        report.exclusions_limitations,
    )
    _add_points_section(
        "D. Denial Reasons (as Stated or Implied by the Insurer)",
        report.denial_reasons,
    )

    # E. Dispute angles
    doc.add_heading("E. Possible Dispute Angles to Explore (Not Legal Advice)", level=1)
    if not report.dispute_angles:
        p = doc.add_paragraph("No dispute angles identified.")
        p.runs[0].italic = True
    else:
        for angle in report.dispute_angles:
            text = angle.text.strip()
            if not text:
                continue
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(text)
            if angle.citations:
                cits = ", ".join(c for c in angle.citations if c.strip())
                if cits:
                    para.add_run(f"\nCitations: {cits}").italic = True

    # F. Missing info
    doc.add_heading("F. Missing Information / Suggested Next Steps", level=1)
    if not report.missing_info:
        p = doc.add_paragraph("No specific missing information identified.")
        p.runs[0].italic = True
    else:
        for item in report.missing_info:
            s = str(item).strip()
            if s:
                doc.add_paragraph(s, style="List Bullet")

    # G. Confidence & verify clauses
    doc.add_heading("G. Confidence and Clauses to Verify", level=1)
    conf = report.confidence
    has_confidence_content = False

    if conf.score is not None:
        para = doc.add_paragraph()
        para.add_run("Confidence score (0–1): ").bold = True
        para.add_run(f"{conf.score:.2f}")
        has_confidence_content = True

    if conf.notes.strip():
        para = doc.add_paragraph()
        para.add_run("Notes: ").bold = True
        para.add_run(conf.notes.strip())
        has_confidence_content = True

    if conf.verify_clauses:
        para = doc.add_paragraph()
        para.add_run("Clauses / sections to double-check:").bold = True
        for clause in conf.verify_clauses:
            s = str(clause).strip()
            if s:
                doc.add_paragraph(s, style="List Bullet")
        has_confidence_content = True

    if not has_confidence_content:
        p = doc.add_paragraph("No explicit confidence metadata provided.")
        p.runs[0].italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


if __name__ == "__main__":
    main()
