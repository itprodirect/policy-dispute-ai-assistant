from __future__ import annotations

import io

from docx import Document

from src.report_builder import render_dispute_docx, render_dispute_markdown
from src.schemas import ConfidenceBlock, DisputeReport, Point


def _sample_report() -> DisputeReport:
    return DisputeReport(
        policy_id="policy-123",
        denial_id="denial-456",
        plain_summary="The dispute summary.",
        coverage_highlights=[Point(text="Coverage may apply.", citation="Coverage A")],
        confidence=ConfidenceBlock(score=0.5, notes="Review source files."),
    )


def test_dispute_markdown_export_includes_human_context() -> None:
    markdown = render_dispute_markdown(
        _sample_report(),
        context={
            "claim_title": "Kitchen water loss",
            "carrier": "Example Carrier",
            "state": "TX",
            "generated_date": "2026-04-26",
            "policy_filename": "policy.pdf",
            "denial_filename": "denial.pdf",
            "demo_mode": True,
        },
    )

    assert "## Report context" in markdown
    assert "- **Claim:** Kitchen water loss" in markdown
    assert "- **Carrier:** Example Carrier" in markdown
    assert "- **State:** TX" in markdown
    assert "- **Generated date:** 2026-04-26" in markdown
    assert "- **Policy file:** policy.pdf" in markdown
    assert "- **Denial file:** denial.pdf" in markdown
    assert "- **Mode:** Demo Mode" in markdown
    assert "AI-generated analysis" in markdown
    assert "not legal advice" in markdown


def test_dispute_markdown_export_handles_no_context() -> None:
    markdown = render_dispute_markdown(_sample_report())

    assert "## Report context" in markdown
    assert "- **Policy ID:** policy-123" in markdown
    assert "- **Denial ID:** denial-456" in markdown
    assert "## A. Plain-English overview of the dispute" in markdown


def test_dispute_markdown_export_includes_analysis_mode() -> None:
    markdown = render_dispute_markdown(
        _sample_report(),
        context={"analysis_mode": "Focused"},
    )

    assert "- **Mode:** Focused" in markdown


def test_dispute_docx_export_includes_human_context() -> None:
    docx_bytes = render_dispute_docx(
        _sample_report(),
        context={
            "claim_title": "Kitchen water loss",
            "state": "TX",
            "generated_date": "2026-04-26",
            "policy_filename": "policy.pdf",
            "denial_filename": "denial.pdf",
        },
    )

    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Report Context" in text
    assert "Claim: Kitchen water loss" in text
    assert "State: TX" in text
    assert "Generated date: 2026-04-26" in text
    assert "Policy file: policy.pdf" in text
    assert "Denial file: denial.pdf" in text


def test_dispute_docx_export_includes_analysis_mode() -> None:
    docx_bytes = render_dispute_docx(
        _sample_report(),
        context={"analysis_mode": "Full"},
    )

    document = Document(io.BytesIO(docx_bytes))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Mode: Full" in text
